from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


path = Path(__file__).resolve().parents[1] / "docs" / "thesis" / "04-thiet-ke-co-so-du-lieu.docx"
errors = []

with ZipFile(path) as archive:
    bad = archive.testzip()
    if bad:
        errors.append(f"Corrupt ZIP member: {bad}")

doc = Document(path)
section = doc.sections[0]
if round(section.page_width.cm, 1) != 21.0 or round(section.page_height.cm, 1) != 29.7:
    errors.append("Page size is not A4")
if len(doc.tables) != 25:
    errors.append(f"Unexpected table count: {len(doc.tables)}")

heading_count = 0
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        heading_count += 1
        if not paragraph.text.strip():
            errors.append("Empty heading")

for index, table in enumerate(doc.tables, start=1):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None or tbl_w.get(qn("w:w")) != "8950":
        errors.append(f"Table {index}: invalid tblW")
    grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
    if sum(grid_widths) != 8950:
        errors.append(f"Table {index}: grid sum {sum(grid_widths)}")
    for rindex, row in enumerate(table.rows):
        if len(row.cells) != len(grid_widths):
            errors.append(f"Table {index} row {rindex}: column mismatch")
        tr_height = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
        if tr_height is not None and tr_height.get(qn("w:hRule")) == "exact":
            errors.append(f"Table {index} row {rindex}: fixed exact height")
        for cindex, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[cindex]:
                errors.append(f"Table {index} row {rindex} cell {cindex}: width mismatch")

print({
    "paragraphs": len(doc.paragraphs),
    "headings": heading_count,
    "tables": len(doc.tables),
    "table_rows": sum(len(table.rows) for table in doc.tables),
    "errors": errors,
})
raise SystemExit(1 if errors else 0)
