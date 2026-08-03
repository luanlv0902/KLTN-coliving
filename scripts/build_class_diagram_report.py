from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_database_report import (
    ACCENT,
    MUTED,
    USABLE_DXA,
    add_body,
    add_bullets,
    add_caption,
    add_heading,
    add_note,
    add_table,
    configure_document,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "thesis" / "05-phan-tich-class-diagram.docx"


PACKAGE_ROWS = [
    ("Identity", "User, PhoneOtp, PasswordResetOtp, AdminLog", "AuthService, UserProfileService, PhoneOtpService, AdminUserService", "Xác thực, hồ sơ, vai trò và quản trị tài khoản"),
    ("Property", "Room, Amenity, RoomImage, RoomVerification, VerificationCheck, ManagerArea", "RoomService, RoomVerificationService, AreaService", "Phòng, tiện ích, bản đồ và kiểm duyệt"),
    ("Rental", "Booking, Contract, ContractEvent, Occupancy, UtilityBill, RentalRoomSnapshot", "BookingService, CapacityService, ContractService, OccupancyService, UtilityBillService", "Thuê phòng, hợp đồng, sức chứa và cư trú"),
    ("Community", "FavoriteRoom, Review, SharedResource, ResourceBooking, SharedSpaceActivity", "FavoriteService, ReviewService, SharedResourceService", "Tương tác cộng đồng và tài nguyên chung"),
    ("Preference", "UserPreference, UserLifestyleProfile, RoomInteraction", "PreferenceService", "Dữ liệu đầu vào cá nhân hóa"),
    ("AI Matching", "UserProfileProjection, RoomProfileProjection, RecommendationResult", "RecommendationService, RoommateMatchingService, ExplanationService", "Tính điểm, xếp hạng và giải thích gợi ý"),
]

STEREOTYPE_ROWS = [
    ("<<entity>>", "Lớp miền có định danh và vòng đời", "User, Room, Booking, Contract, Review"),
    ("<<service>>", "Điều phối quy tắc nghiệp vụ và giao dịch", "BookingService, ContractService, RoomVerificationService"),
    ("<<boundary>>", "Điểm tương tác với actor hoặc service khác", "Web UI, Next.js BFF, internal client"),
    ("<<enumeration>>", "Tập giá trị trạng thái hữu hạn", "Role, RoomStatus, BookingStatus, ContractStatus"),
    ("<<projection>>", "Bản chiếu dữ liệu chỉ chứa trường cần đọc", "RentalRoomSnapshot, UserProfileProjection"),
    ("<<value object>>", "Giá trị không có vòng đời độc lập", "DateRange, CompatibilityScore, ContractSnapshot"),
]

CLASS_GROUPS = {
    "Identity": [
        ("User", "entity", "id, email, password, fullName, phone, role, status", "Thông tin định danh và trạng thái tài khoản"),
        ("PhoneOtp", "entity", "id, userId, phone, codeHash, expiresAt, attemptCount", "Xác thực số điện thoại"),
        ("PasswordResetOtp", "entity", "id, userId, codeHash, expiresAt, consumedAt", "Đặt lại mật khẩu"),
        ("AdminLog", "entity", "adminId, targetUserId, action, oldValue, newValue", "Nhật ký thao tác quản trị"),
        ("AuthService", "service", "-", "register(), login(), getCurrentUser(), updateProfile()"),
        ("PhoneOtpService", "service", "-", "requestOtp(), verifyOtp()"),
        ("AdminUserService", "service", "-", "listUsers(), createUser(), updateUserAction(), getAdminLogs()"),
    ],
    "Property": [
        ("Room", "entity", "id, ownerId, title, address, priceValue, maxOccupants, status", "Thông tin và trạng thái phòng"),
        ("Amenity", "entity", "id, name", "Danh mục tiện ích"),
        ("RoomAmenity", "association", "roomId, amenityId", "Lớp liên kết nhiều-nhiều"),
        ("RoomImage", "entity", "id, roomId, url, sortOrder", "Hình ảnh thuộc phòng"),
        ("RoomVerification", "entity", "roomId, assignedManagerId, recommendation, notes", "Hồ sơ xác minh 1-1 với phòng"),
        ("VerificationCheck", "entity", "verificationId, type, status, checkedById", "Một tiêu chí checklist"),
        ("VerificationDocument", "entity", "verificationId, type, fileUrl, status", "Tài liệu chứng minh"),
        ("CommunityManagerArea", "entity", "managerId, region, city, district, ward", "Phạm vi phụ trách"),
        ("RoomService", "service", "-", "listRooms(), createRoom(), updateRoom(), deleteRoom()"),
        ("RoomVerificationService", "service", "-", "submit(), addDocument(), updateCheck(), managerReview(), adminReview()"),
    ],
    "Rental": [
        ("Booking", "entity", "id, userId, roomId, startDate, endDate, status", "Yêu cầu đặt phòng"),
        ("Contract", "entity", "id, bookingId, hostId, renterId, contentSnapshot, status", "Hợp đồng và vòng đời ký/bàn giao"),
        ("ContractEvent", "entity", "contractId, actorId, type, fromStatus, toStatus", "Lịch sử bất biến của hợp đồng"),
        ("Occupancy", "entity", "roomId, userId, joinedAt, status, terminatedAt", "Trạng thái cư trú thực tế"),
        ("RentalRoomSnapshot", "projection", "roomId, ownerId, status, maxOccupants, currentOccupants", "Bản chiếu phòng phục vụ Rental"),
        ("UtilityBill", "entity", "contractId, month, year, totalCost, proofUrl, status", "Hóa đơn điện nước"),
        ("BookingService", "service", "-", "createBooking(), updateBooking(), cancelBooking(), listHostBookings()"),
        ("CapacityService", "service", "-", "getRoomCapacity(), getRoomsAvailability(), syncRoomOccupancy()"),
        ("ContractService", "service", "-", "create(), sign(), confirmDeposit(), confirmHandover(), renew(), terminate()"),
        ("OccupancyService", "service", "-", "addOccupant(), terminateOccupancy(), occupancyHistory()"),
        ("UtilityBillService", "service", "-", "createBill(), submitProof(), approveBill()"),
    ],
    "Community": [
        ("FavoriteRoom", "entity", "userId, roomId, createdAt", "Phòng được người dùng yêu thích"),
        ("Review", "entity", "userId, roomId, rating, comment, status", "Đánh giá phòng"),
        ("SharedResource", "entity", "roomId, ownerId, type, status, maxDurationMinutes", "Tài nguyên dùng chung"),
        ("ResourceBooking", "entity", "resourceId, userId, startTime, endTime, status", "Lịch đặt tài nguyên"),
        ("SharedSpaceActivity", "entity", "roomId, creatorId, assigneeId, type, eventDate", "Hoạt động cộng đồng"),
        ("ReviewService", "service", "-", "createReview(), updateReview(), listRoomReviews(), updateStatus()"),
        ("SharedResourceService", "service", "-", "createResource(), createBooking(), updateBooking(), createActivity()"),
    ],
    "Preference và AI": [
        ("UserPreference", "entity", "userId, budgetMinVnd, budgetMaxVnd, preferredDistrict", "Tiêu chí chọn phòng"),
        ("UserLifestyleProfile", "entity", "userId, smoking, pets, sleepTime, cleanliness, sociability", "Hồ sơ lối sống"),
        ("RoomInteraction", "entity", "userId, roomId, interactionType, interactionValue", "Tín hiệu hành vi"),
        ("PreferenceService", "service", "-", "getPreference(), upsertPreference(), deletePreference()"),
        ("RecommendationService", "service", "-", "recommendRooms(userId, topK)"),
        ("RoommateMatchingService", "service", "-", "matchRoommates(userId, roomId)"),
        ("ExplanationService", "service", "-", "explainScore(features, score)"),
        ("RecommendationResult", "value object", "roomId, score, explanation", "Kết quả trả về từ AI"),
    ],
}

RELATION_ROWS = [
    ("Room", "RoomImage", "1", "0..*", "Composition", "Hình ảnh không tồn tại độc lập khi phòng bị xóa."),
    ("Room", "RoomVerification", "1", "0..1", "Composition", "Mỗi phòng có tối đa một hồ sơ xác minh."),
    ("RoomVerification", "VerificationCheck", "1", "0..*", "Composition", "Checklist thuộc vòng đời hồ sơ."),
    ("RoomVerification", "VerificationDocument", "1", "0..*", "Composition", "Tài liệu thuộc hồ sơ xác minh."),
    ("Room", "Amenity", "0..*", "0..*", "Association", "Quan hệ qua lớp RoomAmenity."),
    ("Booking", "Contract", "1", "0..1", "Association", "Chỉ booking CONFIRMED mới tạo hợp đồng."),
    ("Contract", "ContractEvent", "1", "0..*", "Composition", "Sự kiện là lịch sử của hợp đồng."),
    ("Contract", "UtilityBill", "1", "0..*", "Composition", "Hóa đơn phụ thuộc hợp đồng."),
    ("SharedResource", "ResourceBooking", "1", "0..*", "Composition", "Lịch đặt phụ thuộc tài nguyên."),
    ("User", "Room/Booking/Contract/Review", "1", "0..*", "Logical association", "Liên service qua ID, không phải FK vật lý."),
]

SERVICE_DEP_ROWS = [
    ("Next.js BFF", "Identity/Property/Rental/Community/Preference/AI", "Gọi API đồng bộ", "Giữ hợp đồng /api cho giao diện và chuyển tiếp identity."),
    ("Property Service", "Identity Service", "Internal client", "Lấy thông tin chủ nhà, reviewer và manager."),
    ("Property Service", "Rental/Community", "Internal client", "Lấy khả dụng, thống kê thuê và thống kê đánh giá."),
    ("Rental Service", "Identity Service", "Internal client", "Ghép hồ sơ ứng viên, người thuê và bên ký hợp đồng."),
    ("Community Service", "Rental Service", "Internal client", "Kiểm tra cư trú ACTIVE và điều kiện đánh giá."),
    ("Community Service", "Property Service", "Internal client", "Kiểm tra phòng và quyền sở hữu tài nguyên."),
    ("AI Matching Service", "Event projections", "Asynchronous dependency", "Đọc bản chiếu, không truy vấn trực tiếp domain database."),
]

ENUM_ROWS = [
    ("Role", "CUSTOMER, HOST, COMMUNITY_MANAGER, ADMIN", "Phân quyền tài khoản; schema còn SERVER/DELIVER để tương thích."),
    ("UserStatus", "ACTIVE, LOCKED, DELETED", "Trạng thái sử dụng tài khoản."),
    ("RoomStatus", "DRAFT, PENDING, NEEDS_REVISION, AVAILABLE, OCCUPIED, REJECTED, HIDDEN", "Vòng đời công bố và kiểm duyệt phòng."),
    ("BookingStatus", "PENDING, CONFIRMED, CANCELLED, COMPLETED", "Vòng đời yêu cầu đặt phòng."),
    ("ContractStatus", "DRAFT, PENDING_RENTER_SIGNATURE, PENDING_DEPOSIT, PENDING_HANDOVER, ACTIVE, EXPIRED, TERMINATED, CANCELLED, DISPUTED", "Vòng đời hợp đồng."),
    ("ReviewStatus", "VISIBLE, HIDDEN, DELETED", "Trạng thái kiểm duyệt đánh giá."),
    ("InteractionType", "VIEW, CLICK, FAVORITE, BOOKING_REQ, CONTRACT", "Tín hiệu hành vi cho AI."),
]

DIAGRAM_ROWS = [
    ("CD-01", "Class Diagram tổng quát", "Các entity trung tâm và quan hệ logic giữa bounded context", "Bắt buộc"),
    ("CD-02", "Identity", "User, OTP, AdminLog và các service xác thực", "Bắt buộc"),
    ("CD-03", "Property", "Room, tiện ích, hình ảnh và quy trình xác minh", "Bắt buộc"),
    ("CD-04", "Rental", "Booking, Contract, Occupancy, UtilityBill và CapacityService", "Bắt buộc"),
    ("CD-05", "Community", "Review, Favorite, SharedResource, ResourceBooking, Activity", "Bắt buộc"),
    ("CD-06", "Preference và AI", "Sở thích, lối sống, tương tác, recommendation và roommate matching", "Bắt buộc"),
    ("CD-07", "Lớp tích hợp/BFF", "Boundary, controller, internal client và dependency giữa service", "Khuyến nghị"),
]


def build():
    doc = Document()
    configure_document(doc, "HỆ THỐNG CO-LIVING NHÀHỢP - PHÂN TÍCH CLASS DIAGRAM")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("PHÂN TÍCH VÀ THIẾT KẾ CLASS DIAGRAM HỆ THỐNG NHÀHỢP")
    set_run_font(run, size=18, bold=True, color=ACCENT)

    add_heading(doc, "1. Mục đích và phạm vi", 1)
    add_body(doc, "Class Diagram mô tả cấu trúc tĩnh của hệ thống thông qua các lớp, thuộc tính, phương thức và quan hệ giữa các lớp. Đối với NhàHợp, sơ đồ lớp được sử dụng để liên kết yêu cầu nghiệp vụ với cấu trúc phần mềm, làm rõ trách nhiệm của từng bounded context và hỗ trợ triển khai, kiểm thử, bảo trì hệ thống.")
    add_body(doc, "Phạm vi phân tích tập trung vào các lớp miền và lớp dịch vụ cốt lõi của Identity, Property, Rental, Community, Preference và AI Matching. Các lớp giao diện chi tiết của React không được đưa đầy đủ vào sơ đồ vì chúng chủ yếu thực hiện trình bày; thay vào đó, Web UI và Next.js BFF được thể hiện như các lớp boundary/controller khi cần mô tả phụ thuộc kiến trúc.")

    add_heading(doc, "2. Quan điểm mô hình hóa", 1)
    add_body(doc, "Mã nguồn backend của đồ án chủ yếu được tổ chức theo module hàm và service thay vì các lớp hướng đối tượng truyền thống. Vì vậy, Class Diagram trong báo cáo là mô hình UML khái niệm ánh xạ từ Prisma model, service module, API contract và quy tắc nghiệp vụ. Sơ đồ không khẳng định rằng mọi phần tử UML đều được khai báo bằng từ khóa class trong mã nguồn.")
    add_note(doc, "Nguyên tắc", "Không đồng nhất bảng cơ sở dữ liệu với lớp phần mềm một cách máy móc. Entity biểu diễn trạng thái miền; Service chứa hành vi nghiệp vụ; Enum biểu diễn trạng thái; Projection và Value Object biểu diễn dữ liệu đọc hoặc giá trị không có vòng đời độc lập.")
    add_caption(doc, "Bảng 5.1. Các stereotype sử dụng trong Class Diagram")
    add_table(doc, ["Stereotype", "Ý nghĩa", "Ví dụ"], STEREOTYPE_ROWS, [1800, 3850, 3300], 10.0)

    add_heading(doc, "3. Phân rã lớp theo bounded context", 1)
    add_body(doc, "Các lớp được nhóm theo ranh giới sở hữu dữ liệu và nghiệp vụ của microservice. Cách phân rã này giúp sơ đồ tổng quát dễ đọc, đồng thời phản ánh đúng kiến trúc database-per-service của hệ thống.")
    add_caption(doc, "Bảng 5.2. Các package lớp của hệ thống")
    add_table(doc, ["Package", "Entity chính", "Service chính", "Trách nhiệm"], PACKAGE_ROWS, [1200, 2600, 2650, 2500], 9.4)

    add_heading(doc, "4. Mô hình lớp tổng quát", 1)
    add_body(doc, "Các lớp trung tâm của NhàHợp gồm User, Room, Booking, Contract, Occupancy, UtilityBill, Review, SharedResource, UserPreference và các projection AI. User tham gia nhiều vai trò nhưng được mô hình hóa bằng một lớp duy nhất kết hợp thuộc tính role, thay vì tạo các lớp con Customer, Host, CommunityManager và Admin.")
    add_body(doc, "Room là gốc tổng hợp của miền Property. Booking và Contract là trung tâm của miền Rental. Contract quản lý nội dung pháp lý, chữ ký, đặt cọc và bàn giao; Occupancy phản ánh trạng thái cư trú thực tế. SharedResource là gốc tổng hợp của lịch tài nguyên. UserPreference và UserLifestyleProfile cung cấp dữ liệu đầu vào cho RecommendationService.")
    add_note(doc, "Lưu ý quan hệ", "Các liên kết User-Room, User-Booking, User-Contract và Room-Review là quan hệ logic giữa service. Khi vẽ sơ đồ tổng quát có thể dùng association nét đứt; khi vẽ sơ đồ lớp vật lý của từng service không biểu diễn chúng như khóa ngoại nội bộ.")

    add_heading(doc, "5. Mô tả các lớp theo package", 1)
    table_no = 3
    for index, (package, rows) in enumerate(CLASS_GROUPS.items(), start=1):
        add_heading(doc, f"5.{index}. Package {package}", 2)
        descriptions = {
            "Identity": "Package Identity quản lý tài khoản, xác thực và nhật ký quản trị. User là entity trung tâm; các thao tác đăng ký, đăng nhập và OTP được đặt trong service để tránh đưa phụ thuộc mã hóa hoặc gửi mã vào entity.",
            "Property": "Package Property quản lý vòng đời phòng từ bản nháp đến khả dụng. RoomVerification cùng checklist và tài liệu tạo thành một aggregate phục vụ quy trình nhân viên quản lý cộng đồng đề nghị và Admin quyết định cuối.",
            "Rental": "Package Rental chứa các lớp cần nhất quán mạnh về thời gian và sức chứa. CapacityService được dùng chung bởi BookingService, ContractService và OccupancyService để tránh vượt maxOccupants.",
            "Community": "Package Community quản lý hành vi xã hội và lịch tài nguyên chung. Service phải kiểm tra quyền cư trú hoặc quyền sở hữu bằng internal client trước khi thay đổi entity.",
            "Preference và AI": "Preference lưu dữ liệu đầu vào do người dùng khai báo và tín hiệu hành vi. AI Matching đọc projection để trả về RecommendationResult; kết quả này là value object, không phải entity nghiệp vụ gốc.",
        }
        add_body(doc, descriptions[package])
        add_caption(doc, f"Bảng 5.{table_no}. Danh mục lớp thuộc package {package}")
        add_table(doc, ["Lớp", "Stereotype", "Thuộc tính chính", "Trách nhiệm/Phương thức"], rows, [1700, 1150, 3000, 3100], 9.2)
        table_no += 1

    add_heading(doc, "6. Quan hệ và bội số giữa các lớp", 1)
    add_body(doc, "Bội số được xác định từ ràng buộc nghiệp vụ và schema đang triển khai. Composition được dùng khi đối tượng con phụ thuộc hoàn toàn vào vòng đời của đối tượng cha; association được dùng khi hai lớp chỉ liên hệ nghiệp vụ. Quan hệ giữa hai microservice được đánh dấu logical association hoặc dependency.")
    add_caption(doc, f"Bảng 5.{table_no}. Quan hệ và bội số quan trọng")
    add_table(doc, ["Lớp A", "Lớp B", "Bội số A", "Bội số B", "Loại quan hệ", "Giải thích"], RELATION_ROWS, [1300, 1550, 850, 850, 1500, 2900], 9.0)
    table_no += 1

    add_heading(doc, "6.1. Composition và aggregation", 2)
    add_body(doc, "RoomImage, RoomVerification, VerificationCheck và VerificationDocument là thành phần phụ thuộc vòng đời của Room hoặc hồ sơ xác minh, do đó phù hợp với composition. ContractEvent và UtilityBill phụ thuộc Contract; ResourceBooking phụ thuộc SharedResource. Khi đối tượng cha bị xóa theo quy tắc cho phép, các đối tượng con tương ứng được xóa cascade trong cùng database.")
    add_body(doc, "Không nên dùng composition cho User với Room hoặc Booking vì dữ liệu này thuộc service khác và cần tồn tại phục vụ hợp đồng, kiểm toán ngay cả khi tài khoản chuyển trạng thái DELETED. Các quan hệ này chỉ là association logic dựa trên ID.")

    add_heading(doc, "6.2. Không sử dụng kế thừa cho vai trò người dùng", 2)
    add_body(doc, "Khách hàng, chủ nhà, nhân viên quản lý cộng đồng và Admin dùng chung cấu trúc tài khoản. Sự khác biệt nằm ở quyền truy cập và hành vi do service kiểm tra. Do đó, mô hình User kết hợp Role enum phù hợp hơn sơ đồ kế thừa User <- Customer/Host/Admin. Cách này cũng phản ánh đúng schema và tránh lặp thuộc tính email, mật khẩu, hồ sơ ở nhiều lớp con.")

    add_heading(doc, "7. Dependency giữa các service", 1)
    add_body(doc, "Class Diagram chi tiết theo package chỉ thể hiện dependency cần thiết thông qua interface hoặc internal client. Service không được truy cập repository/database của service khác. Next.js BFF giữ hợp đồng API dành cho trình duyệt, sau đó gọi service tương ứng với identity đã xác thực.")
    add_caption(doc, f"Bảng 5.{table_no}. Dependency chính giữa các lớp dịch vụ")
    add_table(doc, ["Thành phần nguồn", "Thành phần đích", "Hình thức", "Mục đích"], SERVICE_DEP_ROWS, [1900, 2200, 2000, 2850], 9.4)
    table_no += 1

    add_heading(doc, "8. Các lớp trạng thái và Value Object", 1)
    add_body(doc, "Enum giúp giới hạn trạng thái hợp lệ và làm rõ các nhánh xử lý trong Activity/Sequence Diagram. Khi vẽ, enum được biểu diễn bằng stereotype <<enumeration>> và dependency từ entity sử dụng enum. Không cần nối enum bằng association có bội số.")
    add_caption(doc, f"Bảng 5.{table_no}. Các enumeration quan trọng")
    add_table(doc, ["Enumeration", "Giá trị chính", "Vai trò"], ENUM_ROWS, [1700, 4400, 2850], 9.2)
    table_no += 1
    add_body(doc, "Các value object nên thể hiện gồm DateRange cho khoảng thuê, ContractSnapshot cho nội dung pháp lý đã chốt, CompatibilityScore cho điểm tương thích và RecommendationResult cho kết quả gợi ý. Chúng không có khóa định danh riêng và thường được tạo lại từ dữ liệu đầu vào.")

    add_heading(doc, "9. Ánh xạ từ Class Diagram sang mã nguồn", 1)
    mapping_rows = [
        ("Entity", "Prisma model trong schema của service", "Room -> property schema; Contract -> rental schema"),
        ("Service", "Module .cjs/.ts xuất các hàm nghiệp vụ", "ContractService -> services/rental-service/contracts.cjs"),
        ("Boundary/Controller", "Next.js route hoặc Express endpoint", "/api/contracts/* -> /v1/contracts/*"),
        ("Repository", "Prisma Client được truyền vào service", "prisma.booking, prisma.contract, prisma.room"),
        ("Internal client", "Module gọi API giữa bounded context", "identityClient, domainClients, rental-client"),
        ("Projection", "Bảng/DTO bản chiếu", "RentalRoomSnapshot, ai.room_profiles"),
    ]
    add_caption(doc, f"Bảng 5.{table_no}. Quy tắc ánh xạ UML sang triển khai")
    add_table(doc, ["Phần tử UML", "Thành phần mã nguồn", "Ví dụ"], mapping_rows, [1750, 3650, 3550], 9.5)
    table_no += 1

    add_heading(doc, "10. Tính nhất quán với các sơ đồ hành vi", 1)
    add_bullets(doc, [
        "Activity đăng ký và đăng nhập được hiện thực bởi Web UI, BFF, AuthService và User.",
        "Activity đặt phòng tương ứng Customer -> BookingService -> CapacityService -> Booking.",
        "Quy trình duyệt phòng sử dụng RoomVerificationService cùng RoomVerification, VerificationCheck và các enum trạng thái.",
        "Sequence hợp đồng ánh xạ trực tiếp vào ContractService, Contract, ContractEvent, Booking và Occupancy.",
        "Sequence tài nguyên chung sử dụng SharedResourceService, SharedResource, ResourceBooking cùng dependency đến Rental/Property.",
        "Luồng AI sử dụng PreferenceService, UserPreference, RoomInteraction, RecommendationService và RecommendationResult.",
    ])

    add_heading(doc, "11. Danh mục Class Diagram cần trình bày", 1)
    add_body(doc, "Một sơ đồ duy nhất chứa toàn bộ entity và service sẽ quá rộng, nhiều đường giao nhau và khó đặt trong trang báo cáo. Nên sử dụng một sơ đồ tổng quát để trình bày khái niệm, sau đó tách theo bounded context để thể hiện thuộc tính và phương thức.")
    add_caption(doc, f"Bảng 5.{table_no}. Danh sách Class Diagram đề xuất")
    add_table(doc, ["Mã", "Tên sơ đồ", "Phạm vi", "Mức độ"], DIAGRAM_ROWS, [1000, 2450, 4050, 1450], 9.6)

    add_heading(doc, "12. Quy tắc trình bày sơ đồ", 1)
    add_bullets(doc, [
        "Mỗi class chỉ giữ thuộc tính và phương thức quan trọng; không đưa toàn bộ trường audit vào sơ đồ tổng quát.",
        "Entity, Service, Enum, Projection và Value Object phải có stereotype rõ ràng.",
        "Ghi bội số ở hai đầu association và dùng hình thoi đặc chỉ cho composition thực sự.",
        "Dependency giữa service dùng mũi tên nét đứt; không dùng association như quan hệ đối tượng nội bộ.",
        "Không vẽ khóa ngoại cơ sở dữ liệu như một thuộc tính lặp nếu association đã thể hiện rõ, trừ khi ID liên service là chủ đích của thiết kế.",
        "Tên lớp dùng PascalCase, thuộc tính và phương thức dùng camelCase; enum viết tên giá trị in hoa.",
        "Sơ đồ tổng quát ưu tiên khả năng đọc; sơ đồ package mới trình bày đầy đủ hơn thuộc tính và operation.",
    ])

    add_heading(doc, "13. Kết luận", 1)
    add_body(doc, "Class Diagram của NhàHợp phản ánh cấu trúc miền phân tán theo microservice. Entity giữ trạng thái nghiệp vụ, service điều phối quy tắc và giao dịch, enum kiểm soát vòng đời, còn projection/value object hỗ trợ đọc dữ liệu và AI. Việc tách sơ đồ theo bounded context giúp mô hình bám sát mã nguồn, hạn chế phụ thuộc chéo và thống nhất với Use Case, Activity Diagram, Sequence Diagram và thiết kế cơ sở dữ liệu đã xây dựng.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
