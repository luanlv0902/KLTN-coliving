from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "thesis" / "04-thiet-ke-co-so-du-lieu.docx"

# Narrative proposal preset with a named academic A4 override for a Vietnamese thesis.
FONT = "Times New Roman"
INK = "1F1F1F"
MUTED = "666666"
HEADER_FILL = "D9E2F3"
SUBHEADER_FILL = "EEF3F8"
ACCENT = "1F4E79"
USABLE_DXA = 8950


def set_run_font(run, size=13, bold=None, italic=None, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == USABLE_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "110")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=11, italic=True, color=MUTED)
    return p


def add_table(doc, headers, rows, widths, font_size=10.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, bold=True, color="17365D")

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, SUBHEADER_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=11.5, bold=True, color=ACCENT)
    r = p.add_run(text)
    set_run_font(r, size=11.5)
    set_table_geometry(table, [USABLE_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc, header_text="HỆ THỐNG CO-LIVING NHÀHỢP - THIẾT KẾ CƠ SỞ DỮ LIỆU"):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    heading_tokens = {
        "Heading 1": (16, 14, 7),
        "Heading 2": (14, 11, 5),
        "Heading 3": (13, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    set_run_font(hr, size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


SERVICE_ROWS = [
    ("1", "Identity", "IDENTITY_DATABASE_URL", "User, OTP, nhật ký quản trị, sự kiện định danh"),
    ("2", "Property", "PROPERTY_DATABASE_URL", "Phòng, tiện ích, hình ảnh, hồ sơ xác minh và phân vùng"),
    ("3", "Rental", "RENTAL_DATABASE_URL", "Đặt phòng, hợp đồng, cư trú, hóa đơn và projection phòng"),
    ("4", "Community", "COMMUNITY_DATABASE_URL", "Yêu thích, đánh giá, tài nguyên chung và hoạt động cộng đồng"),
    ("5", "Preference", "PREFERENCE_DATABASE_URL", "Sở thích, lối sống và tương tác phục vụ cá nhân hóa"),
    ("6", "AI projection", "AI_USE_PROJECTIONS", "Projection người dùng, phòng, cư trú và tương tác cho mô hình AI"),
]

CATALOGS = {
    "Identity": [
        ("User", "Tài khoản và hồ sơ định danh", "id", "email UNIQUE"),
        ("PhoneOtp", "Mã OTP xác thực số điện thoại", "id", "FK userId -> User.id"),
        ("PasswordResetOtp", "OTP đặt lại mật khẩu", "id", "FK userId -> User.id"),
        ("AdminLog", "Nhật ký thao tác quản trị", "id", "FK adminId/targetUserId -> User.id"),
        ("IdentityInboxEvent", "Chống xử lý lặp sự kiện đến", "id", "eventId UNIQUE"),
        ("IdentityOutboxEvent", "Hàng đợi sự kiện định danh", "id", "Index trạng thái và thời điểm thử lại"),
    ],
    "Property": [
        ("Room", "Thông tin phòng cho thuê", "id", "Index ownerId, status"),
        ("Amenity", "Danh mục tiện ích", "id", "name UNIQUE"),
        ("RoomAmenity", "Liên kết phòng - tiện ích", "id", "UNIQUE(roomId, amenityId)"),
        ("RoomImage", "Hình ảnh phòng", "id", "FK roomId -> Room.id"),
        ("RoomVerification", "Hồ sơ xác minh phòng", "id", "roomId UNIQUE"),
        ("VerificationCheck", "Các mục checklist xác minh", "id", "UNIQUE(verificationId, type)"),
        ("RoomVerificationDocument", "Tài liệu xác minh", "id", "FK verificationId"),
        ("CommunityManagerArea", "Khu vực nhân viên phụ trách", "id", "UNIQUE theo phạm vi địa lý"),
        ("PropertyOutboxEvent", "Hàng đợi sự kiện Property", "id", "Index trạng thái và aggregate"),
    ],
    "Rental": [
        ("Booking", "Yêu cầu đặt phòng", "id", "Index userId, roomId, status"),
        ("Contract", "Hợp đồng thuê", "id", "contractNumber và bookingId UNIQUE"),
        ("ContractEvent", "Lịch sử vòng đời hợp đồng", "id", "FK contractId -> Contract.id"),
        ("Occupancy", "Trạng thái cư trú thực tế", "id", "UNIQUE(roomId, userId)"),
        ("RentalRoomSnapshot", "Projection phòng phục vụ Rental", "roomId", "Không FK sang Property"),
        ("UtilityBill", "Hóa đơn điện nước", "id", "UNIQUE(contractId, month, year)"),
        ("Invoice", "Chứng từ gắn với booking", "id", "bookingId UNIQUE"),
        ("Payment", "Thanh toán của Invoice", "id", "invoiceId UNIQUE"),
        ("RentalOutboxEvent", "Hàng đợi sự kiện Rental", "id", "Index trạng thái và aggregate"),
    ],
    "Community": [
        ("UserDeviceToken", "Token thiết bị nhận thông báo", "id", "token UNIQUE"),
        ("FavoriteRoom", "Phòng người dùng yêu thích", "id", "UNIQUE(userId, roomId)"),
        ("Review", "Đánh giá phòng", "id", "UNIQUE(roomId, userId)"),
        ("SharedResource", "Tài nguyên dùng chung", "id", "Index roomId và ownerId"),
        ("ResourceBooking", "Lịch đặt tài nguyên", "id", "FK resourceId -> SharedResource.id"),
        ("SharedSpaceActivity", "Hoạt động trong không gian chung", "id", "Index roomId, creatorId, assigneeId"),
        ("CommunityOutboxEvent", "Hàng đợi sự kiện Community", "id", "Index trạng thái và aggregate"),
    ],
    "Preference": [
        ("user_preferences", "Tiêu chí ưu tiên tìm phòng", "id", "userId UNIQUE"),
        ("user_lifestyle_profiles", "Hồ sơ thói quen sống", "id", "userId UNIQUE"),
        ("RoomInteraction", "Tương tác người dùng - phòng", "id", "Index userId và roomId"),
        ("PreferenceOutboxEvent", "Hàng đợi sự kiện Preference", "id", "Index trạng thái và aggregate"),
    ],
    "AI": [
        ("ai.user_profiles", "Projection người dùng và sở thích", "user_id", "Dữ liệu đọc tối ưu cho AI"),
        ("ai.room_profiles", "Projection đặc trưng phòng", "room_id", "Dữ liệu đọc tối ưu cho AI"),
        ("ai.occupancy_profiles", "Projection thành viên đang cư trú", "room_id + user_id", "Index room_id, status"),
        ("ai.room_interactions", "Projection tương tác huấn luyện/xếp hạng", "interaction_id", "Index user_id và room_id"),
        ("ai.processed_events", "Sự kiện đã tiêu thụ", "event_id", "Chống xử lý lặp"),
        ("ai.projection_reconciliation_runs", "Lịch sử đồng bộ projection", "id", "Theo dõi trạng thái và lỗi"),
    ],
}


CORE_DICTIONARIES = [
    ("User", [
        ("id", "String/UUID", "Khóa chính của người dùng."),
        ("email", "String", "Email đăng nhập; duy nhất trong Identity DB."),
        ("password", "String", "Mật khẩu đã băm; không trả về cho service khác."),
        ("fullName, name", "String", "Tên đầy đủ và tên hiển thị."),
        ("phone, phoneVerified, phoneVerifiedAt", "String/Boolean/DateTime", "Số điện thoại và trạng thái xác thực OTP."),
        ("role", "Role", "CUSTOMER, HOST, COMMUNITY_MANAGER hoặc ADMIN; schema còn giá trị tương thích cũ."),
        ("status", "UserStatus", "ACTIVE, LOCKED hoặc DELETED."),
        ("address, avatarUrl, birthDate, gender", "Nullable", "Thông tin hồ sơ cá nhân không bắt buộc."),
        ("latitude, longitude", "Float?", "Tọa độ tùy chọn của người dùng."),
        ("createdAt, updatedAt", "DateTime", "Thời điểm tạo và cập nhật gần nhất."),
    ]),
    ("Room", [
        ("id", "String/UUID", "Khóa chính của phòng."),
        ("ownerId", "String?", "ID logic của chủ nhà trong Identity Service."),
        ("title, description, address", "String", "Thông tin mô tả bắt buộc."),
        ("city, district, ward và các mã", "String?", "Cấu trúc địa chỉ phục vụ lọc và phân vùng."),
        ("latitude, longitude", "Float?", "Tọa độ hiển thị trên bản đồ."),
        ("priceValue, priceText", "BigInt?/String?", "Giá trị số dùng tính toán và chuỗi dùng hiển thị."),
        ("areaValue, areaText", "Decimal?/String?", "Diện tích chuẩn hóa và nội dung hiển thị."),
        ("maxOccupants, currentOccupants", "Int?", "Sức chứa tối đa và số người đang ở."),
        ("allowPets, allowSmoking", "Boolean?", "Chính sách vật nuôi và hút thuốc."),
        ("cleanlinessRequired, noiseTolerance, guestPolicy, preferredSleepHabit", "String?", "Các thuộc tính tương thích lối sống."),
        ("status", "RoomStatus", "DRAFT, PENDING, NEEDS_REVISION, AVAILABLE, OCCUPIED, REJECTED hoặc HIDDEN."),
        ("createdAt, updatedAt", "DateTime", "Thời điểm tạo và cập nhật."),
    ]),
    ("RoomVerification", [
        ("id, roomId", "UUID", "Khóa chính và quan hệ 1-1 với Room."),
        ("reviewerId, assignedManagerId", "String?", "ID logic của Admin và nhân viên được phân công."),
        ("submittedAt, managerAssignedAt, managerReviewedAt, reviewedAt", "DateTime?", "Các mốc xử lý hồ sơ."),
        ("managerRecommendation", "Enum", "PENDING, NEEDS_REVISION, RECOMMEND_APPROVAL hoặc RECOMMEND_REJECTION."),
        ("revisionReason, rejectionReason, adminNote, managerNote", "String?", "Ghi chú và lý do xử lý."),
        ("identityPassed ... legalOccupancyPassed", "Boolean", "Kết quả tổng hợp các tiêu chí xác minh."),
        ("informationAccurateConfirmed ...", "Boolean", "Ba cam kết pháp lý bắt buộc của chủ nhà."),
        ("declarationAcceptedAt, declarationVersion, IP, userAgent", "Nullable", "Bằng chứng chấp thuận cam kết."),
    ]),
    ("Booking", [
        ("id", "String/UUID", "Khóa chính yêu cầu đặt phòng."),
        ("userId, roomId", "String", "ID logic của khách hàng và phòng."),
        ("startDate, endDate", "DateTime", "Khoảng thời gian thuê được yêu cầu."),
        ("status", "BookingStatus", "PENDING, CONFIRMED, CANCELLED hoặc COMPLETED."),
        ("cancelledAt, cancelledById, cancellationActor, cancellationReason", "Nullable", "Thông tin truy vết thao tác hủy."),
        ("createdAt, updatedAt", "DateTime", "Thời điểm tạo và cập nhật."),
    ]),
    ("Contract", [
        ("id, contractNumber", "UUID/String", "Khóa chính và số hợp đồng duy nhất."),
        ("bookingId", "String? UNIQUE", "Liên kết 0..1 với một Booking đã CONFIRMED."),
        ("roomId, hostId, renterId", "String", "ID logic của phòng và hai bên tham gia."),
        ("startDate, endDate, monthlyRent", "DateTime/Float", "Thời hạn và giá thuê hàng tháng."),
        ("depositAmount, depositStatus, depositPaidAt", "Float/Enum/DateTime?", "Thông tin tiền đặt cọc."),
        ("paymentDueDay, paymentMethod", "Int/String?", "Ngày đến hạn và phương thức thanh toán."),
        ("electricityRate, waterRate, utilitiesNotes", "Float?/String?", "Đơn giá và ghi chú điện nước."),
        ("contentSnapshot, contentHash, termsVersion", "Json/String", "Nội dung pháp lý bất biến và mã kiểm tra toàn vẹn."),
        ("hostSignedAt ... renterSignatureUserAgent", "Nullable", "Bằng chứng chữ ký điện tử của hai bên."),
        ("hostHandoverConfirmedAt, renterHandoverConfirmedAt", "DateTime?", "Xác nhận bàn giao hai chiều."),
        ("status", "ContractStatus", "Vòng đời từ DRAFT đến ACTIVE, EXPIRED, TERMINATED..."),
        ("terminatedAt, terminationReason, renewalCount", "Nullable/Int", "Dữ liệu gia hạn và chấm dứt."),
    ]),
    ("Occupancy", [
        ("id", "String/UUID", "Khóa chính bản ghi cư trú."),
        ("roomId, userId", "String", "Cặp duy nhất xác định một người trong một phòng."),
        ("joinedAt", "DateTime", "Thời điểm bắt đầu cư trú."),
        ("status", "String", "ACTIVE hoặc INACTIVE theo luồng hiện tại."),
        ("terminatedAt, terminationReason", "Nullable", "Thời điểm và lý do kết thúc cư trú."),
        ("notes", "String?", "Ghi chú quản lý."),
    ]),
    ("UtilityBill", [
        ("id, contractId", "UUID", "Khóa chính và FK đến Contract."),
        ("month, year", "Int", "Kỳ hóa đơn; duy nhất trong phạm vi hợp đồng."),
        ("previousReading, currentReading", "Float?", "Chỉ số kỳ trước và kỳ hiện tại."),
        ("electricityUsage, waterUsage", "Float?", "Lượng điện và nước tiêu thụ."),
        ("electricityCost, waterCost, totalCost", "Float", "Chi phí thành phần và tổng chi phí."),
        ("status", "String", "PENDING hoặc PAID theo luồng hiện tại."),
        ("paymentProofUrl, paymentProofSubmittedAt", "Nullable", "Minh chứng thanh toán của người thuê."),
        ("approvedAt", "DateTime?", "Thời điểm chủ nhà xác nhận thanh toán."),
    ]),
    ("SharedResource", [
        ("id, roomId, ownerId", "String", "Khóa chính và ID logic của phòng/chủ sở hữu."),
        ("name, description, type", "String", "Tên, mô tả và loại EQUIPMENT/SPACE."),
        ("status", "String", "ACTIVE, MAINTENANCE hoặc BUSY theo xử lý vận hành."),
        ("requiresApproval", "Boolean", "Cấu hình tài nguyên có yêu cầu duyệt."),
        ("maxDurationMinutes", "Int", "Thời lượng tối đa cho một lần đặt."),
        ("createdAt, updatedAt", "DateTime", "Thời điểm tạo và cập nhật."),
    ]),
    ("ResourceBooking", [
        ("id, resourceId", "String", "Khóa chính và FK đến SharedResource."),
        ("userId", "String", "ID logic của người đang cư trú đặt tài nguyên."),
        ("title", "String?", "Mục đích sử dụng."),
        ("startTime, endTime", "DateTime", "Khoảng thời gian đặt."),
        ("status", "String", "PENDING, APPROVED hoặc CANCELLED trong luồng nghiệp vụ."),
        ("createdAt, updatedAt", "DateTime", "Thời điểm tạo và cập nhật."),
    ]),
    ("Review", [
        ("id, userId, roomId", "String", "Khóa chính và ID logic của người dùng/phòng."),
        ("rating", "Int", "Điểm đánh giá; API kiểm tra miền giá trị hợp lệ."),
        ("comment", "String?", "Nội dung nhận xét."),
        ("status", "ReviewStatus", "VISIBLE, HIDDEN hoặc DELETED."),
        ("roomId + userId", "UNIQUE", "Mỗi người chỉ có một đánh giá trên một phòng."),
    ]),
    ("user_preferences", [
        ("id, userId", "UUID/String", "Khóa chính và userId duy nhất."),
        ("budgetMinVnd, budgetMaxVnd", "BigInt?", "Khoảng ngân sách thuê phòng."),
        ("preferredDistrict", "String?", "Khu vực ưu tiên."),
        ("lifestyleArchetype", "String?", "Nhóm phong cách sống tổng hợp."),
        ("priorityCleanliness, prioritySocialEnvironment", "Int?", "Trọng số ưu tiên mặc định bằng 3."),
        ("acceptSmokingRoommates, acceptPets", "Boolean?", "Khả năng chấp nhận hút thuốc và vật nuôi."),
    ]),
    ("user_lifestyle_profiles", [
        ("id, userId", "UUID/String", "Khóa chính và userId duy nhất."),
        ("smoking, pets, petTolerance", "Enum/Boolean", "Thói quen hút thuốc và vật nuôi."),
        ("sleepTime, wakeupTime", "Enum/Int?", "Khung giờ ngủ và thức dậy."),
        ("cleanliness, noiseTolerance, sociability, privacyLevel", "Int?", "Các chỉ số lối sống."),
        ("guestsFrequency, cookingFrequency", "Enum/Int?", "Tần suất khách đến và nấu ăn."),
        ("budget", "BigInt?", "Ngân sách tham khảo cho ghép phòng/bạn cùng phòng."),
    ]),
]


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("THIẾT KẾ CƠ SỞ DỮ LIỆU HỆ THỐNG NHÀHỢP")
    set_run_font(run, size=18, bold=True, color=ACCENT)

    add_heading(doc, "1. Mục tiêu và phạm vi thiết kế", 1)
    add_body(doc, "Cơ sở dữ liệu của NhàHợp được thiết kế để lưu trữ toàn bộ vòng đời nghiệp vụ co-living, từ định danh người dùng, công bố và kiểm duyệt phòng, đặt phòng, ký hợp đồng, quản lý cư trú đến vận hành tài nguyên dùng chung và cá nhân hóa bằng AI. Thiết kế phải bảo đảm tính nhất quán của dữ liệu nghiệp vụ, khả năng truy vết các thay đổi quan trọng và khả năng mở rộng độc lập của từng nhóm chức năng.")
    add_body(doc, "Khác với mô hình cơ sở dữ liệu nguyên khối, phiên bản hiện tại áp dụng nguyên tắc database-per-service. Mỗi microservice sở hữu schema và dữ liệu thuộc bounded context của mình. Các giá trị như userId hoặc roomId xuất hiện ở nhiều service là tham chiếu logic; chúng không được triển khai thành khóa ngoại xuyên cơ sở dữ liệu.")
    add_note(doc, "Phạm vi", "Phần này mô tả schema đang được sử dụng trong kiến trúc microservice. Schema Prisma nguyên khối ở thư mục gốc chỉ còn vai trò tương thích cho chế độ fallback và không phải mô hình vật lý đích để vẽ ERD triển khai.")

    add_heading(doc, "2. Công nghệ và kiến trúc lưu trữ", 1)
    add_body(doc, "Hệ thống sử dụng PostgreSQL làm hệ quản trị cơ sở dữ liệu quan hệ. Prisma ORM được dùng tại các dịch vụ Node.js để khai báo model, migration và truy vấn dữ liệu có kiểu. Riêng AI Matching Service sử dụng các bảng projection trong schema ai, được tối ưu cho thao tác đọc, tính đặc trưng và xếp hạng.")
    add_caption(doc, "Bảng 4.1. Phân vùng sở hữu dữ liệu của hệ thống NhàHợp")
    add_table(doc, ["STT", "Vùng dữ liệu", "Cấu hình", "Trách nhiệm"], SERVICE_ROWS, [600, 1450, 2100, 4800], 10.2)

    add_heading(doc, "2.1. Nguyên tắc database-per-service", 2)
    add_bullets(doc, [
        "Mỗi service chỉ ghi trực tiếp vào database do chính service đó sở hữu.",
        "Không tạo khóa ngoại vật lý giữa các database; việc xác minh tham chiếu được thực hiện qua API nội bộ hoặc projection.",
        "Dữ liệu cần tính nhất quán mạnh được đặt trong cùng bounded context. Booking, Contract và Occupancy cùng thuộc Rental Service để kiểm tra sức chứa trong một giao dịch.",
        "Các thay đổi liên dịch vụ được truyền bằng sự kiện và Transactional Outbox, tránh cập nhật kép không nguyên tử.",
        "AI chỉ đọc các projection cần thiết, không truy vấn trực tiếp toàn bộ bảng nghiệp vụ của các service.",
    ])

    add_heading(doc, "2.2. Chuẩn dữ liệu chung", 2)
    add_body(doc, "Khóa chính của phần lớn bảng sử dụng UUID dưới dạng String để có thể sinh độc lập tại từng service. Các bảng đều ưu tiên cặp createdAt và updatedAt để phục vụ truy vết. Trạng thái nghiệp vụ quan trọng được khai báo bằng enum; các trạng thái linh hoạt hơn trong module cộng đồng và hóa đơn hiện được lưu dưới dạng String nhưng vẫn được giới hạn ở tầng service.")
    add_body(doc, "Các chỉ mục được đặt trên trường thường xuyên dùng để tìm kiếm hoặc phân quyền như userId, ownerId, roomId, status, assignedManagerId và thời điểm sự kiện. Ràng buộc UNIQUE được sử dụng cho email, số hợp đồng, cặp phòng-tiện ích, cặp phòng-người cư trú và kỳ hóa đơn theo hợp đồng.")

    add_heading(doc, "3. Danh mục bảng dữ liệu", 1)
    table_no = 2
    for service, rows in CATALOGS.items():
        add_heading(doc, f"3.{list(CATALOGS.keys()).index(service) + 1}. {service} database", 2)
        add_caption(doc, f"Bảng 4.{table_no}. Danh mục bảng thuộc {service}")
        add_table(doc, ["Tên bảng", "Chức năng", "Khóa chính", "Ràng buộc/Chỉ mục chính"], rows, [1900, 3000, 1250, 2800], 9.8)
        table_no += 1

    add_heading(doc, "4. Thiết kế chi tiết các vùng dữ liệu", 1)
    add_heading(doc, "4.1. Identity database", 2)
    add_body(doc, "Identity Service là nguồn dữ liệu chuẩn duy nhất của thông tin định danh. Bảng User lưu thông tin đăng nhập, hồ sơ, vai trò và trạng thái tài khoản. PhoneOtp và PasswordResetOtp lưu mã băm, thời hạn và số lần thử thay vì lưu OTP ở dạng rõ. AdminLog lưu vết hành động quản trị; inbox/outbox hỗ trợ trao đổi sự kiện an toàn.")

    add_heading(doc, "4.2. Property database", 2)
    add_body(doc, "Property Service quản lý nội dung phòng và quy trình xác minh. Room có quan hệ vật lý với RoomImage, RoomAmenity và RoomVerification. Một Room chỉ có tối đa một RoomVerification; mỗi hồ sơ xác minh có nhiều tài liệu và nhiều mục checklist. CommunityManagerArea không có FK sang User vì tài khoản nhân viên thuộc Identity Service, nhưng managerId được xác minh thông qua hợp đồng API nội bộ.")

    add_heading(doc, "4.3. Rental database", 2)
    add_body(doc, "Rental Service tập trung các dữ liệu cần nhất quán giao dịch. Booking có thể phát sinh tối đa một Contract. Khi hợp đồng được hai bên xác nhận bàn giao, hệ thống cập nhật Contract thành ACTIVE, tạo hoặc kích hoạt Occupancy và chuyển Booking thành COMPLETED trong cùng giao dịch. UtilityBill phụ thuộc Contract và được giới hạn một hóa đơn cho mỗi tháng/năm của một hợp đồng.")
    add_body(doc, "RentalRoomSnapshot là projection cục bộ của Room. Bảng chỉ chứa các trường Rental cần để kiểm tra quyền sở hữu, trạng thái, sức chứa và hiển thị booking/hợp đồng. Cách tiếp cận này loại bỏ việc Rental truy vấn trực tiếp Property DB, đồng thời cho phép kiểm tra sức chứa bằng giao dịch tuần tự hóa.")

    add_heading(doc, "4.4. Community database", 2)
    add_body(doc, "Community Service lưu các tương tác xã hội và vận hành không gian chung. FavoriteRoom và Review giữ userId, roomId dưới dạng ID logic. SharedResource có quan hệ 1-n với ResourceBooking; khi tạo lịch, service gọi Rental để kiểm tra người dùng có Occupancy ACTIVE và gọi Property khi cần xác nhận quyền sở hữu của chủ nhà.")

    add_heading(doc, "4.5. Preference database", 2)
    add_body(doc, "Preference Service sở hữu dữ liệu sở thích và lối sống. user_preferences chứa tiêu chí chọn phòng có tính tường minh như ngân sách và khu vực. user_lifestyle_profiles chứa thói quen dùng để tính độ tương thích. RoomInteraction ghi nhận tín hiệu hành vi VIEW, CLICK, FAVORITE, BOOKING_REQ và CONTRACT để phục vụ xếp hạng.")

    add_heading(doc, "4.6. AI projection schema", 2)
    add_body(doc, "AI Matching Service không sở hữu dữ liệu nghiệp vụ gốc. Các bảng ai.user_profiles, ai.room_profiles, ai.occupancy_profiles và ai.room_interactions là bản chiếu được cập nhật từ sự kiện của các service. processed_events bảo đảm một sự kiện không bị áp dụng lặp; projection_reconciliation_runs ghi nhận các lần đối soát và khôi phục projection.")

    add_heading(doc, "5. Từ điển dữ liệu các bảng cốt lõi", 1)
    add_body(doc, "Các bảng dưới đây tập trung vào trường có ý nghĩa trực tiếp đối với nghiệp vụ và các ràng buộc quan trọng. Các trường thời gian chuẩn createdAt/updatedAt được lược bớt ở những bảng mà ý nghĩa đã rõ để tránh lặp lại.")
    for table_name, rows in CORE_DICTIONARIES:
        add_heading(doc, f"5.{CORE_DICTIONARIES.index((table_name, rows)) + 1}. Bảng {table_name}", 2)
        add_caption(doc, f"Bảng 4.{table_no}. Từ điển dữ liệu bảng {table_name}")
        add_table(doc, ["Trường/Nhóm trường", "Kiểu dữ liệu", "Ý nghĩa và ràng buộc"], rows, [2500, 1750, 4700], 9.7)
        table_no += 1

    add_heading(doc, "6. Quan hệ dữ liệu", 1)
    add_heading(doc, "6.1. Quan hệ vật lý trong từng database", 2)
    physical_rows = [
        ("Identity", "User 1-n PhoneOtp; User 1-n PasswordResetOtp; User 1-n AdminLog theo hai vai trò admin/target."),
        ("Property", "Room n-n Amenity qua RoomAmenity; Room 1-n RoomImage; Room 1-1 RoomVerification; RoomVerification 1-n Document và Check."),
        ("Rental", "Booking 0..1-1 Contract; Booking 0..1-1 Invoice; Invoice 1-0..1 Payment; Contract 1-n ContractEvent và UtilityBill."),
        ("Community", "SharedResource 1-n ResourceBooking. Các bảng Review, FavoriteRoom và Activity chỉ lưu ID logic của domain khác."),
        ("Preference", "Các bảng tách độc lập theo userId; không tạo FK vật lý sang Identity để giữ ranh giới service."),
    ]
    add_caption(doc, f"Bảng 4.{table_no}. Quan hệ khóa ngoại vật lý")
    add_table(doc, ["Database", "Quan hệ chính"], physical_rows, [1600, 7350], 10.2)
    table_no += 1

    add_heading(doc, "6.2. Quan hệ logic liên dịch vụ", 2)
    logical_rows = [
        ("User.id", "Room.ownerId", "Identity -> Property", "Xác định chủ sở hữu phòng."),
        ("User.id", "Booking.userId", "Identity -> Rental", "Xác định khách hàng gửi yêu cầu."),
        ("User.id", "Contract.hostId/renterId", "Identity -> Rental", "Xác định hai bên hợp đồng; nội dung ký được chụp vào contentSnapshot."),
        ("Room.id", "Booking.roomId / Contract.roomId", "Property -> Rental", "Liên kết nghiệp vụ qua roomId và RentalRoomSnapshot."),
        ("Room.id", "Review.roomId / SharedResource.roomId", "Property -> Community", "Liên kết đánh giá và tài nguyên với phòng."),
        ("User.id", "user_preferences.userId", "Identity -> Preference", "Gắn hồ sơ cá nhân hóa với người dùng."),
        ("Các ID nguồn", "ai.*_profiles", "Services -> AI", "Đồng bộ bằng sự kiện; không phải dữ liệu gốc."),
    ]
    add_caption(doc, f"Bảng 4.{table_no}. Quan hệ logic giữa các bounded context")
    add_table(doc, ["ID nguồn", "Trường tham chiếu", "Hướng dữ liệu", "Ý nghĩa"], logical_rows, [1600, 2250, 2050, 3050], 9.8)
    table_no += 1

    add_note(doc, "Quy tắc vẽ ERD", "ERD vật lý chỉ nối các bảng có khóa ngoại trong cùng database. Quan hệ liên service nên được thể hiện ở sơ đồ dữ liệu logic bằng đường nét đứt và ghi chú 'logical reference' hoặc 'projection', không dùng ký hiệu FK vật lý.")

    add_heading(doc, "7. Quản lý trạng thái nghiệp vụ", 1)
    state_rows = [
        ("Room", "DRAFT -> PENDING -> AVAILABLE", "Có thể rẽ sang NEEDS_REVISION, REJECTED, HIDDEN; AVAILABLE có thể chuyển OCCUPIED theo projection cư trú."),
        ("Booking", "PENDING -> CONFIRMED -> COMPLETED", "Có thể chuyển CANCELLED khi chủ nhà từ chối hoặc khách hàng hủy đúng điều kiện."),
        ("Contract", "DRAFT -> PENDING_RENTER_SIGNATURE -> PENDING_DEPOSIT/PENDING_HANDOVER -> ACTIVE", "Sau ACTIVE có thể EXPIRED, TERMINATED hoặc DISPUTED; trạng thái chữ ký chủ nhà được chấp nhận từ DRAFT."),
        ("Occupancy", "ACTIVE -> INACTIVE", "Được kích hoạt khi bàn giao hoàn tất hoặc chủ nhà thêm người hợp lệ."),
        ("ResourceBooking", "PENDING -> APPROVED/CANCELLED", "Tạo mới ở PENDING trong luồng microservice hiện tại."),
        ("UtilityBill", "PENDING -> PAID", "Chỉ chuyển PAID sau khi chủ nhà duyệt minh chứng."),
        ("Review", "VISIBLE -> HIDDEN/DELETED", "Admin có thể ẩn, khôi phục hoặc xóa logic."),
    ]
    add_caption(doc, f"Bảng 4.{table_no}. Các vòng đời trạng thái chính")
    add_table(doc, ["Đối tượng", "Luồng trạng thái chính", "Nhánh/Điều kiện"], state_rows, [1500, 3150, 4300], 9.8)
    table_no += 1

    add_heading(doc, "8. Bảo đảm toàn vẹn và nhất quán", 1)
    add_heading(doc, "8.1. Ràng buộc và chỉ mục", 2)
    add_bullets(doc, [
        "Ràng buộc UNIQUE ngăn dữ liệu trùng: email, contractNumber, bookingId của Contract, kỳ UtilityBill, cặp Room-Amenity, Room-User trong Occupancy và Room-User trong Review.",
        "Khóa ngoại trong cùng database sử dụng onDelete: Cascade ở dữ liệu phụ thuộc chặt như OTP, hình ảnh, checklist, tài liệu, sự kiện hợp đồng và hóa đơn.",
        "Chỉ mục status hỗ trợ các màn hình hàng đợi như phòng chờ duyệt, booking chờ xử lý, hợp đồng đang hiệu lực và hóa đơn chưa thanh toán.",
        "Chỉ mục userId, roomId và ownerId hỗ trợ phân quyền, tra cứu hồ sơ và tổng hợp thống kê.",
    ])

    add_heading(doc, "8.2. Giao dịch và kiểm soát cạnh tranh", 2)
    add_body(doc, "Các thao tác nhạy cảm với sức chứa như xác nhận booking, bàn giao hợp đồng và thêm người cư trú chạy trong giao dịch tuần tự hóa. Khi đặt tài nguyên chung, Community Service khóa bản ghi tài nguyên trước khi kiểm tra các khoảng thời gian PENDING/APPROVED để hạn chế hai người đặt trùng cùng một khung giờ.")
    add_body(doc, "Hợp đồng lưu contentSnapshot và contentHash tại thời điểm tạo. Nhờ đó, nội dung đã ký không bị thay đổi khi người dùng cập nhật hồ sơ hoặc phòng thay đổi thông tin hiển thị. ContractEvent ghi lại mỗi bước quan trọng cùng actorId, trạng thái trước/sau, địa chỉ IP và user agent.")

    add_heading(doc, "8.3. Transactional Outbox và projection", 2)
    add_body(doc, "Identity, Property, Rental, Community và Preference đều có bảng OutboxEvent với trạng thái, số lần thử, thời điểm thử lại, lỗi cuối cùng và thời điểm phát hành. Thay đổi domain và bản ghi outbox được ghi trong cùng giao dịch. Worker phát sự kiện lên RabbitMQ; nếu broker lỗi, sự kiện vẫn còn trong database để thử lại.")
    add_body(doc, "Sau khi Occupancy thay đổi, Rental ghi RentalOutboxEvent. Property tiêu thụ sự kiện để cập nhật currentOccupants và trạng thái phòng. AI tiêu thụ các sự kiện để cập nhật projection, đồng thời dùng processed_events để bảo đảm idempotency. Cơ chế đối soát định kỳ giúp tái tạo projection khi phát hiện thiếu hoặc lệch dữ liệu.")

    add_heading(doc, "9. Bảo mật và dữ liệu cá nhân", 1)
    add_bullets(doc, [
        "Mật khẩu và OTP chỉ lưu dạng băm; token đăng nhập không được lưu trong bảng User.",
        "Mỗi service chỉ nhận các trường định danh tối thiểu qua API nội bộ; password không được cung cấp cho Property, Rental, Community, Preference hoặc AI.",
        "Các endpoint service nội bộ yêu cầu internal service token và truyền userId/role đã xác thực từ BFF.",
        "Tài khoản xóa được đánh dấu DELETED và làm mờ thông tin cá nhân thay vì xóa cứng, nhằm giữ toàn vẹn tham chiếu hợp đồng và nhật ký.",
        "URL minh chứng thanh toán, tài liệu xác minh, chữ ký, IP và user agent là dữ liệu nhạy cảm; quyền đọc phải giới hạn theo vai trò và quan hệ với đối tượng nghiệp vụ.",
        "AdminLog, ContractEvent và các bảng inbox/outbox cung cấp khả năng kiểm toán nhưng cần chính sách lưu trữ và dọn dữ liệu theo thời hạn.",
    ])

    add_heading(doc, "10. Các sơ đồ dữ liệu cần trình bày", 1)
    add_body(doc, "Do kiến trúc phân tán, phần minh họa cơ sở dữ liệu nên gồm một sơ đồ tổng quan dữ liệu logic và các ERD vật lý theo bounded context. Không nên ghép toàn bộ 41 bảng vào một ERD duy nhất vì sơ đồ sẽ khó đọc và thể hiện sai ranh giới sở hữu dữ liệu.")
    erd_rows = [
        ("ERD-01", "Sơ đồ dữ liệu logic toàn hệ thống", "User, Room, Booking, Contract, Occupancy, Community, Preference và AI projection."),
        ("ERD-02", "Identity database", "User, OTP, AdminLog, Inbox/Outbox."),
        ("ERD-03", "Property database", "Room, Amenity, Image, Verification, Check, Document, ManagerArea."),
        ("ERD-04", "Rental database", "Booking, Contract, ContractEvent, Occupancy, Snapshot, UtilityBill, Invoice, Payment."),
        ("ERD-05", "Community database", "Favorite, Review, SharedResource, ResourceBooking, Activity, DeviceToken."),
        ("ERD-06", "Preference và AI projection", "Sở thích, lối sống, tương tác và các bảng projection AI."),
    ]
    add_caption(doc, f"Bảng 4.{table_no}. Danh mục sơ đồ cơ sở dữ liệu đề xuất")
    add_table(doc, ["Mã", "Tên sơ đồ", "Phạm vi"], erd_rows, [1100, 2850, 5000], 10.0)

    add_heading(doc, "11. Kết luận", 1)
    add_body(doc, "Thiết kế cơ sở dữ liệu của NhàHợp phản ánh trực tiếp kiến trúc microservice và vòng đời nghiệp vụ co-living. Việc tách dữ liệu theo bounded context giúp các module định danh, phòng, thuê ở, cộng đồng, sở thích và AI có thể phát triển độc lập. Tính nhất quán được duy trì bằng ràng buộc trong từng database, giao dịch cho các thao tác quan trọng, snapshot/projection cho dữ liệu đọc liên miền và Transactional Outbox cho sự kiện liên dịch vụ.")
    add_body(doc, "Mô hình này đáp ứng yêu cầu hiện tại của đồ án, đồng thời tạo nền tảng để mở rộng số lượng phòng, người dùng và chức năng cá nhân hóa mà không phụ thuộc vào một schema nguyên khối.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
