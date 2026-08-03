from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def extract(path: Path) -> dict:
    document = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = clean(paragraph.text)
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
            }
        )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    sections = []
    for paragraph in paragraphs:
        style = paragraph["style"].lower()
        text = paragraph["text"]
        if style.startswith("heading") or (
            len(text) <= 180
            and (
                text.upper().startswith("CHƯƠNG ")
                or text[:1].isdigit()
                or text.upper() in {"MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO"}
            )
        ):
            sections.append(paragraph)

    return {
        "path": str(path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "section_candidates": sections,
    }


def main() -> None:
    if len(sys.argv) < 3:
        raise SystemExit("Usage: extract_docx_structure.py INPUT.docx OUTPUT.json")
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(extract(input_path), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
